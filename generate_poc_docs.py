from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


ROOT = Path(r"C:\codex02")
OUT = ROOT / "artifacts"
OUT.mkdir(exist_ok=True)

TODAY = date.today().isoformat()

IMAGE_DIR = Path(
    r"C:\Users\Admin\OneDrive\文档\xwechat_files\dreamshope_0893\temp\RWTemp\2026-06\9e20f478899dc29eb19741386f9343c8"
)
SOURCE_IMAGES = [
    ("图1 Azure构成概念图", IMAGE_DIR / "af01e416f3f6934956161382251e0eb3.jpg"),
    ("图2 Azure资源清单", IMAGE_DIR / "72bd11cbe8a5d56e9a5220fac33098bb.jpg"),
    ("图3 ServiceNow验证范围", IMAGE_DIR / "e636edceecb803c703a48a3e6c6f9b2d.jpg"),
]

DOCX_PATH = OUT / "ServiceNow_Discovery_Azure_PoC_构筑手册.docx"
XLSX_PATH = OUT / "ServiceNow_Discovery_Azure_PoC_资源清单与检查表.xlsx"


SOURCES = [
    {
        "id": "SN1",
        "title": "ServiceNow Discovery overview",
        "url": "https://www.servicenow.com/docs/r/it-operations-management/discovery/discovery.html",
        "use": "Discovery 概念、PoC范围定义",
    },
    {
        "id": "SN2",
        "title": "ServiceNow Discovery phases",
        "url": "https://www.servicenow.com/docs/r/it-operations-management/discovery/c_DiscoProcessFlows.html",
        "use": "Shazzam / Classification / Identification / Exploration 阶段说明",
    },
    {
        "id": "SN3",
        "title": "ServiceNow MID Server",
        "url": "https://www.servicenow.com/docs/r/servicenow-platform/mid-server/mid-server.html",
        "use": "MID Server 安装、验证和运行前提",
    },
    {
        "id": "SN4",
        "title": "Data collected during discovery of Windows computers",
        "url": "https://www.servicenow.com/docs/r/it-operations-management/itom-visibility/r_DataCollDiscoWindowsComputers.html",
        "use": "Windows Discovery 验证项目参考",
    },
    {
        "id": "SN5",
        "title": "Data collected during discovery of Linux computers",
        "url": "https://www.servicenow.com/docs/r/it-operations-management/itom-visibility/r_DataCollDiscoLinuxComputers.html",
        "use": "Linux Discovery 验证项目参考",
    },
    {
        "id": "SN6",
        "title": "Cloud Discovery collected data",
        "url": "https://www.servicenow.com/docs/r/it-operations-management/itom-visibility/cloud-discovery-collected-data.html",
        "use": "Azure Cloud Discovery 采集对象参考",
    },
    {
        "id": "SN7",
        "title": "Service Mapping",
        "url": "https://www.servicenow.com/docs/r/it-operations-management/service-mapping/service-mapping.html",
        "use": "Top-down Entry Point / Application Service 验证参考",
    },
    {
        "id": "SN8",
        "title": "Available Discovery and Service Mapping Patterns",
        "url": "https://www.servicenow.com/docs/r/it-operations-management/discovery-and-service-mapping-patterns/available-patterns.html",
        "use": "OS / Network device / Service Mapping pattern 可用性检查",
    },
    {
        "id": "SN9",
        "title": "Agent Client Collector landing page",
        "url": "https://www.servicenow.com/docs/r/it-operations-management/agent-client-collector/acc-landing-page.html",
        "use": "ACC 验证入口",
    },
    {
        "id": "SN10",
        "title": "ACC data collection",
        "url": "https://www.servicenow.com/docs/r/it-operations-management/agent-client-collector/acc-data-collection.html",
        "use": "ACC Host Data / Check / 采集内容参考",
    },
    {
        "id": "AZ1",
        "title": "Virtual network peering overview",
        "url": "https://learn.microsoft.com/en-us/azure/virtual-network/virtual-network-peering-overview",
        "use": "VNet Peering / 私有IP互通",
    },
    {
        "id": "AZ2",
        "title": "Azure Bastion configuration settings",
        "url": "https://learn.microsoft.com/en-us/azure/bastion/configuration-settings",
        "use": "AzureBastionSubnet 命名和 /26 以上要求",
    },
    {
        "id": "AZ3",
        "title": "Create an Azure Bastion host",
        "url": "https://learn.microsoft.com/en-us/azure/bastion/tutorial-create-host-portal",
        "use": "Bastion Portal 创建步骤",
    },
    {
        "id": "AZ4",
        "title": "Quickstart: Create a NAT gateway",
        "url": "https://learn.microsoft.com/en-us/azure/nat-gateway/quickstart-create-nat-gateway-portal",
        "use": "NAT Gateway 创建与子网关联",
    },
    {
        "id": "AZ5",
        "title": "Network security groups overview",
        "url": "https://learn.microsoft.com/en-us/azure/virtual-network/network-security-groups-overview",
        "use": "NSG 入站/出站规则设计",
    },
    {
        "id": "AZ6",
        "title": "Private IP addresses in Azure",
        "url": "https://learn.microsoft.com/en-us/azure/virtual-network/ip-services/private-ip-addresses",
        "use": "Azure 子网保留 IP，避免手动使用 .1/.2/.3",
    },
    {
        "id": "AZ7",
        "title": "Quickstart: Create a Windows VM in the Azure portal",
        "url": "https://learn.microsoft.com/en-us/azure/virtual-machines/windows/quick-create-portal",
        "use": "Windows VM Portal 创建步骤",
    },
    {
        "id": "AZ8",
        "title": "Quickstart: Create a Linux VM in the Azure portal",
        "url": "https://learn.microsoft.com/en-us/azure/virtual-machines/linux/quick-create-portal",
        "use": "Linux VM Portal 创建步骤",
    },
    {
        "id": "AZ9",
        "title": "Create and manage Windows VMs with multiple NICs",
        "url": "https://learn.microsoft.com/en-us/azure/virtual-machines/windows/multiple-nics",
        "use": "多 NIC 只能挂到同一 VNet 内多个 subnet",
    },
    {
        "id": "AZ10",
        "title": "Create an application and service principal",
        "url": "https://learn.microsoft.com/en-us/entra/identity-platform/howto-create-service-principal-portal",
        "use": "Azure Cloud Discovery 用 App registration / client secret",
    },
    {
        "id": "AZ11",
        "title": "Assign Azure roles using the Azure portal",
        "url": "https://learn.microsoft.com/en-us/azure/role-based-access-control/role-assignments-portal",
        "use": "给 Service Principal 分配 Reader 角色",
    },
    {
        "id": "AZ12",
        "title": "Azure Virtual Network FAQ",
        "url": "https://learn.microsoft.com/en-us/azure/virtual-network/virtual-networks-faq",
        "use": "VNet peering 不支持 transitive peering；vnet02 与 vnet03 需要直接 peering",
    },
]


VALIDATION_SCOPE = [
    ["Horizontal", "IP Range", "Credential-less", "Windows / Linux", "对象", "用于验证无凭据扫描、端口发现、分类前置行为；不会证明完整属性采集。"],
    ["Horizontal", "IP Range", "Credential", "Windows / Linux / NW(PaloAlto) / NW(FortiGate)", "对象", "Windows/Linux 使用 OS 凭据，网络设备使用 SNMP/设备凭据。NW(Others) 不做。"],
    ["Horizontal", "Cloud Discovery", "Credential", "Azure", "对象", "AWS / Others Cloud 不做。通过 Azure Service Principal 做 API 级 Cloud Discovery。"],
    ["Horizontal + Service Mapping", "IP Range", "Credential", "Windows / Linux / NW(PaloAlto) / NW(FortiGate)", "对象", "先完成 Horizontal Discovery，再验证依赖关系和基础连接。"],
    ["Horizontal + Service Mapping", "IP Range + Top-down Entry Point", "Credential", "Windows / Linux / NW(PaloAlto) / NW(FortiGate)", "对象", "需要准备 IIS/Nginx 等可访问入口点；NW 只作为链路/依赖对象验证。"],
    ["Horizontal + Service Mapping", "Cloud Discovery + Top-down Entry Point", "Credential", "Azure", "对象", "Azure 资源发现后，以云上应用入口点做映射验证。"],
    ["Horizontal + Service Mapping", "Cloud Discovery + traffic-based Mapping", "Credential", "検証不可", "对象外", "按图中范围，不在本 PoC 做。"],
    ["ACC", "Agent Based", "Credential", "Windows / Linux", "对象", "在 Windows 11 与 Linux/Ubuntu 主机上安装 ACC Agent，验证 Host Data / Check 结果。"],
]


RESOURCE_PLAN = [
    ["管理", "Resource Group", "RG-SNOW-Discovery-POC", "1", "Japan East", "图2指定名称；按客户要求改区域。"],
    ["网络", "Virtual Network", "vnet01-server-network", "1", "10.0.0.0/24", "Server & Network Device Zone"],
    ["网络", "Virtual Network", "vnet02-mid-linux", "1", "172.16.0.0/24", "MID & Linux Discovery Zone"],
    ["网络", "Virtual Network", "vnet03-client-acc", "1", "192.168.0.0/24", "Client & ACC Zone"],
    ["网络", "Subnet", "subnet-server", "1", "10.0.0.0/25", "vnet01 workload subnet。图中 .1/.2/.3 仅作概念。"],
    ["网络", "Subnet", "AzureBastionSubnet", "1", "10.0.0.128/26", "Bastion 专用；必须使用此名称。"],
    ["网络", "Subnet", "subnet-mid", "1", "172.16.0.0/25", "Linux MID / Ubuntu 22 / VNet02 访问测试 VM。"],
    ["网络", "Subnet", "subnet-client", "1", "192.168.0.0/24", "Windows 11 / ACC client zone。"],
    ["网络", "Network Security Group", "nsg-vnet01-server", "1", "关联 subnet-server", "只允许 MID 私有 IP 到目标端口。"],
    ["网络", "Network Security Group", "nsg-vnet02-mid", "1", "关联 subnet-mid", "允许管理入口、出站 HTTPS、Discovery 流量。"],
    ["网络", "Network Security Group", "nsg-vnet03-client", "1", "关联 subnet-client", "允许 MID 到 ACC/OS 验证所需端口。"],
    ["网络", "VNet Peering", "peer-vnet01-vnet02 / peer-vnet01-vnet03 / peer-vnet02-vnet03", "3组/6方向", "Full mesh peering。Portal 中确认双向。"],
    ["网络", "NAT Gateway", "natgw-vnet01", "1", "subnet-server", "让无公网 IP VM 出站访问 ServiceNow/更新源。"],
    ["网络", "NAT Gateway", "natgw-vnet02", "1", "subnet-mid", "图1/2对应 NAT Gateway 02。"],
    ["网络", "NAT Gateway", "natgw-vnet03", "0-1", "subnet-client", "图中未列，但 ACC/Windows 更新需要出站时建议追加。"],
    ["远程接続", "Azure Bastion", "bas-snow-poc", "0-1", "vnet01", "正式评审或不想给 VM 公网 IP 时推荐。"],
    ["监视", "Log Analytics Workspace", "law-snow-poc", "1", "任意", "诊断日志/网络排错用。"],
    ["Compute", "Windows MID Server VM", "vm-win-mid-01", "1", "D2s_v5 / Windows Server 2022", "安装 Windows MID Server。"],
    ["Compute", "Linux MID Server VM", "vm-linux-mid-01", "1", "B2s / Ubuntu 22.04", "安装 Linux MID Server。"],
    ["Compute", "Windows Server 2019 VM", "vm-win2019-01", "1", "B2ms", "Windows Discovery / IIS service mapping target。"],
    ["Compute", "Windows Server 2022 VM", "vm-win2022-01", "1", "B2ms", "Windows Discovery target。"],
    ["Compute", "Windows 11 VM", "vm-win11-acc-01", "1", "B2ms", "ACC Windows target；注意 Windows client license/image 可用性。"],
    ["Compute", "Ubuntu 22 VM", "vm-ubuntu-01", "1", "B2ms", "Linux Discovery / Nginx target；同时用于验证 vnet02↔vnet03 连通性。"],
    ["Compute", "RHEL 9 VM", "vm-rhel9-01", "1", "B2ms", "Linux Discovery target。"],
    ["Network Device", "Palo Alto VM-Series", "vm-pa-01", "1", "Marketplace / PAYG", "Discovery対象のみ；启用管理接口/SNMP。"],
    ["Network Device", "FortiGate VM", "vm-fgt-01", "1", "Marketplace / PAYG 1 vCPU 推奨", "Discovery対象のみ；启用管理接口/SNMP。"],
]


IP_PLAN = [
    ["区域", "VNet", "Subnet", "CIDR", "资源", "建议私有 IP", "说明"],
    ["Server & Network Device", "vnet01", "subnet-server", "10.0.0.0/25", "vm-win-mid-01", "10.0.0.10", "Windows MID"],
    ["Server & Network Device", "vnet01", "subnet-server", "10.0.0.0/25", "vm-win2019-01", "10.0.0.20", "Windows Server 2019"],
    ["Server & Network Device", "vnet01", "subnet-server", "10.0.0.0/25", "vm-win2022-01", "10.0.0.21", "Windows Server 2022"],
    ["Server & Network Device", "vnet01", "subnet-server", "10.0.0.0/25", "vm-rhel9-01", "10.0.0.30", "RHEL 9"],
    ["Server & Network Device", "vnet01", "subnet-server", "10.0.0.0/25", "vm-pa-01", "10.0.0.40", "Palo Alto management"],
    ["Server & Network Device", "vnet01", "subnet-server", "10.0.0.0/25", "vm-fgt-01", "10.0.0.41", "FortiGate management"],
    ["MID & Linux", "vnet02", "subnet-mid", "172.16.0.0/25", "vm-linux-mid-01", "172.16.0.10", "Linux MID"],
    ["MID & Linux", "vnet02", "subnet-mid", "172.16.0.0/25", "vm-ubuntu-01", "172.16.0.20", "Ubuntu 22 / vnet02→vnet03 连通性测试"],
    ["Client & ACC", "vnet03", "subnet-client", "192.168.0.0/24", "vm-win11-acc-01", "192.168.0.10", "Windows 11 + ACC"],
]


NSG_RULES = [
    ["NSG", "方向", "优先级", "名称", "来源", "目标", "协议/端口", "Action", "用途"],
    ["nsg-vnet01-server", "Inbound", "100", "Allow-MID-to-Targets", "10.0.0.10,172.16.0.10", "10.0.0.0/25", "TCP/UDP Any（初期）", "Allow", "先跑通 Discovery；验收后可收紧。"],
    ["nsg-vnet01-server", "Inbound", "110", "Allow-SNMP-from-MID", "10.0.0.10,172.16.0.10", "10.0.0.40-10.0.0.41", "UDP 161", "Allow", "网络设备 Discovery。"],
    ["nsg-vnet01-server", "Inbound", "120", "Allow-Windows-Discovery", "10.0.0.10,172.16.0.10", "Windows targets", "TCP 135,445,5985,5986 + RPC", "Allow", "Windows WMI/WinRM/SMB。"],
    ["nsg-vnet01-server", "Inbound", "130", "Allow-Linux-SSH", "10.0.0.10,172.16.0.10", "Linux targets", "TCP 22", "Allow", "Linux/RHEL SSH。"],
    ["nsg-vnet02-mid", "Outbound", "100", "Allow-MID-HTTPS-Out", "MID VMs", "Internet", "TCP 443", "Allow", "MID 到 ServiceNow instance / Azure API。"],
    ["nsg-vnet02-mid", "Inbound", "100", "Allow-VNet-Peering", "VirtualNetwork", "VirtualNetwork", "Any", "Allow", "Peering 内部互通。"],
    ["nsg-vnet03-client", "Inbound", "100", "Allow-Ubuntu-VNet02-Access", "172.16.0.0/24", "vm-win11-acc-01", "TCP 5985", "Allow", "用于验证 Ubuntu 从 vnet02 访问 vnet03 的私网连通性。"],
    ["nsg-vnet03-client", "Outbound", "100", "Allow-ACC-HTTPS-Out", "vm-win11-acc-01", "Internet", "TCP 443", "Allow", "ACC Agent / Windows Update / ServiceNow 通信。"],
    ["全部", "Inbound", "4096", "Deny-Internet-Direct", "Internet", "Any", "Any", "Deny", "不建议给 VM 公网 RDP/SSH；用 Bastion。"],
]


BUILD_STEPS = [
    ["0-1", "准备 Azure/ServiceNow 前提", "确认 Azure subscription、Owner/Contributor 权限、ServiceNow 管理员/Discovery admin/MID 用户、ITOM Discovery/Service Mapping/ACC 许可或插件。", "能登录 Azure Portal 和 ServiceNow；能创建 RG/VM/Network。"],
    ["0-2", "决定区域和命名", "默认使用 Japan East；如客户指定区域则统一替换。资源名前缀使用 snow-poc。", "所有资源写入资源清单。"],
    ["1-1", "创建 Resource Group", "Azure Portal > Resource groups > Create：RG-SNOW-Discovery-POC。", "RG 创建成功。"],
    ["1-2", "创建 vnet01", "Virtual networks > Create：10.0.0.0/24；添加 subnet-server 10.0.0.0/25 和 AzureBastionSubnet 10.0.0.128/26。", "两个 subnet 存在；AzureBastionSubnet 名称精确。"],
    ["1-3", "创建 vnet02", "Virtual networks > Create：172.16.0.0/24；subnet-mid 172.16.0.0/25。Ubuntu 与 Linux MID 都放在这里。", "vnet02 只有一个工作 subnet；不要再创建跨 VNet 的第二 NIC subnet。"],
    ["1-4", "创建 vnet03", "Virtual networks > Create：192.168.0.0/24；subnet-client 192.168.0.0/24。", "Client/ACC subnet 存在。"],
    ["1-5", "创建 NSG 并关联 subnet", "创建 nsg-vnet01-server/nsg-vnet02-mid/nsg-vnet03-client，按 NSG_Rules 添加初期规则。", "NSG 关联到对应 subnet。"],
    ["1-6", "创建 NAT Gateway", "按图创建 natgw-vnet01、natgw-vnet02 并关联 subnet-server/subnet-mid；若 vnet03 无公网 IP 且需 ACC/更新，追加 natgw-vnet03。", "VM 不带公网 IP 也可出站访问 Internet。"],
    ["1-7", "创建 Full Mesh VNet Peering", "vnet01-vnet02、vnet01-vnet03、vnet02-vnet03。Portal 中确认 reciprocal/bidirectional。", "三对 VNet 互相 Connected。"],
    ["1-8", "创建 Azure Bastion", "Bastion > Create：放在 vnet01；选择 AzureBastionSubnet；创建 Public IP。", "可从 Portal 用 Bastion 登录 vnet01/vnet02/vnet03 VM。"],
    ["2-1", "创建 Windows MID VM", "Windows Server 2022 / D2s_v5 / subnet-server / 私有 IP 10.0.0.10 / 不开放公网 RDP。", "Bastion 能登录；出站 443 可访问 ServiceNow。"],
    ["2-2", "创建 Linux MID VM", "Ubuntu 22.04 / B2s / subnet-mid / 私有 IP 172.16.0.10 / SSH Key。", "Bastion/Serial/SSH 可登录；出站 443 正常。"],
    ["2-3", "创建 Discovery targets", "创建 Windows 2019/2022、RHEL9、Ubuntu22、Windows11。Windows11 如 Marketplace license 不满足，先用 Windows Server 作为 ACC 替代验证并记录偏差。", "所有 VM Running；私有 IP 符合 IP_Plan。"],
    ["2-4", "验证 Ubuntu 同时访问 VNet02 / VNet03", "Ubuntu 22 保持单 NIC 放在 vnet02/subnet-mid；确认通过 vnet02↔vnet03 peering 能访问 vnet03 的验证目标（建议 vm-win11-acc-01 的 5985 端口或你实际开放的测试端口）。", "Ubuntu 只看到一块网卡，但能通过私有 IP 访问 vnet03 目标。"],
    ["2-5", "部署 Palo Alto / FortiGate", "从 Azure Marketplace 部署 PAYG。目标只是被 Discovery 发现，优先启用管理接口、SNMP read-only、HTTPS/SSH 管理。", "从 MID 私有 IP 可 ping/SSH/HTTPS/SNMP 到设备管理 IP。"],
    ["2-6", "OS 基础配置", "Windows 开启远程管理/WMI/防火墙例外；Linux 开启 SSH；在 vnet03 的验证目标上开放 5985 或另一个测试端口；测试 Nmap/credential-less 所需端口；安装 IIS/Nginx 作为 Service Mapping entry point。", "MID 到各目标端口连通；Ubuntu 可到达 vnet03 验证端口。"],
    ["3-1", "安装 Windows MID Server", "ServiceNow instance 中下载 MID Server 安装包，在 vm-win-mid-01 安装，使用专用 MID 用户，启动服务并在 ServiceNow 中 Validate。", "MID 状态 Up / Validated。"],
    ["3-2", "安装 Linux MID Server", "在 vm-linux-mid-01 下载 Linux MID 包，配置 wrapper/服务，启动并 Validate。", "Linux MID 状态 Up / Validated。"],
    ["3-3", "创建 Credentials", "ServiceNow 中创建 Windows、SSH、SNMP、Azure Service Principal、ACC 相关凭据。不要把密码写入手册或 Excel。", "Credentials 测试通过或在 Discovery Log 中能看到成功使用。"],
    ["3-4", "创建 Discovery Schedules", "按 vnet01/vnet02/vnet03 的 CIDR 建立 credential-less 与 credentialed schedules；指定对应 MID 或 MID selection。", "Discovery Status 运行完成，无重大 credential error。"],
    ["3-5", "配置 Azure Cloud Discovery", "Azure Entra ID 创建 App registration/client secret；订阅上给 Reader 权限；ServiceNow 创建 Azure Credential/Cloud account/schedule。", "能发现 Azure VM/VNet/Subnet/NIC/Disk 等云资源。"],
    ["3-6", "配置 Service Mapping", "准备 IIS/Nginx 入口点；创建 Application Service/Entry Point；运行 Top-down Discovery。", "Application Service 中出现入口点、主机和依赖关系。"],
    ["3-7", "配置 ACC", "在 Windows11 和 Linux 目标安装 ACC Agent，使用 registration key/installer；关联 Check Policy。", "ACC Agent Online；Host Data/Checks 有结果。"],
    ["4-1", "验证与截图", "按 Verification_Matrix 逐项截图：Azure资源、MID状态、Discovery Status、CI、Cloud资源、Service Mapping、ACC结果，以及 Ubuntu→VNet03 连通性。", "证据包完整。"],
    ["5-1", "成本控制/清理", "PoC 每日结束停止 VM；结束后删除 RG 或按资源清单逐项清理 Marketplace/NAT/Bastion/Public IP。", "无遗留计费资源。"],
]


SERVICENOW_SETUP = [
    ["项目", "ServiceNow侧配置", "Azure/OS侧动作", "成功标准", "证据"],
    ["MID Server", "创建专用 MID 用户，下载并安装 Windows/Linux MID，Validate。", "VM 出站 TCP 443 到 instance；时间同步；DNS 可解析。", "MID Server Up / Validated。", "MID Server 列表截图。"],
    ["Windows Credential Discovery", "Windows Credential；Discovery Schedule 包含 Windows IP。", "Windows 防火墙允许 WMI/SMB/WinRM；本地管理员或授权账号。", "发现 Windows Server CI，属性/软件/服务有采集。", "Discovery Status + CI 详情。"],
    ["Linux Credential Discovery", "SSH Credential；Discovery Schedule 包含 Linux IP。", "SSH 22 开启；账号可执行基础命令；必要时 sudo。", "发现 Linux/RHEL/Ubuntu CI。", "Discovery Status + CI 详情。"],
    ["Credential-less", "IP range schedule 使用无凭据路径或先做 Shazzam/端口发现。", "NSG 允许来自 MID 的扫描流量。", "能看到 IP/开放端口/分类前置结果；不要求完整 CI 属性。", "Discovery log / status。"],
    ["Network Device", "SNMP Credential / Network device Discovery。", "Palo Alto/FortiGate 管理口启用 SNMP read-only，NSG 允许 UDP161。", "发现网络设备 CI；厂商/型号/SNMP属性可见。", "CI + SNMP 测试记录。"],
    ["Azure Cloud Discovery", "Azure Credential / Cloud Account / Schedule。", "App registration + client secret；Subscription Reader role。", "发现 Azure VM/VNet/Subnet/NIC/Disk/Resource Group。", "Cloud Discovery status + CI/Cloud resources。"],
    ["Service Mapping", "Application Service + Entry Point；运行 top-down mapping。", "IIS/Nginx 可访问；目标 OS 已被发现；端口开放。", "Application Service Map 出现入口点、主机、依赖。", "Service Map 截图。"],
    ["ACC", "Registration Key / Check Policy / ACC-V content。", "Windows11/Linux 安装 agent；出站 443。", "Agent Online；Host Data/Check 有结果。", "ACC Agent/Host Data 页面。"],
]


VERIFICATION_MATRIX = [
    ["ID", "验证范围", "对象", "前提", "执行步骤", "期待结果", "证据", "判定"],
    ["V-01", "Horizontal / IP Range / Credential-less", "Windows, Linux", "MID Up；NSG 放通扫描", "运行 credential-less IP range schedule。", "发现 IP、开放端口、可能的分类结果；不会期待完整属性。", "Discovery Status/Log", ""],
    ["V-02", "Horizontal / IP Range / Credential", "Windows Server 2019/2022", "Windows credential 可用", "运行 Windows credentialed Discovery。", "生成/更新 Windows Server CI，关键属性采集。", "CI详情/Discovery Log", ""],
    ["V-03", "Horizontal / IP Range / Credential", "Ubuntu/RHEL", "SSH credential 可用", "运行 Linux credentialed Discovery。", "生成 Linux Server CI，OS/CPU/内存/磁盘/网卡等属性可见。", "CI详情/Discovery Log", ""],
    ["V-04", "Horizontal / IP Range / Credential", "Palo Alto", "SNMP/管理口可达", "运行 Network Discovery。", "发现 Palo Alto 设备 CI；厂商/型号/SNMP属性可见。", "CI详情/SNMP测试", ""],
    ["V-05", "Horizontal / IP Range / Credential", "FortiGate", "SNMP/管理口可达", "运行 Network Discovery。", "发现 FortiGate 设备 CI；厂商/型号/SNMP属性可见。", "CI详情/SNMP测试", ""],
    ["V-06", "Horizontal / Cloud Discovery", "Azure", "Service Principal + Reader role", "运行 Azure Cloud Discovery schedule。", "发现 Azure Resource Group/VNet/VM/NIC/Disk 等云资源。", "Cloud Discovery status", ""],
    ["V-07", "Service Mapping / IP Range", "Windows/Linux app", "IIS/Nginx入口可访问", "创建 Entry Point 并运行 top-down discovery。", "Application Service Map 生成，入口点与主机关系可见。", "Service Map 截图", ""],
    ["V-08", "Service Mapping / Cloud + Entry Point", "Azure-hosted app", "Cloud Discovery + app入口", "用 Azure VM 上的 Web 入口运行 top-down。", "云资源与应用服务能建立可解释关系。", "Service Map + Cloud CI", ""],
    ["V-09", "ACC / Agent Based", "Windows 11", "ACC Agent安装/注册", "安装 ACC Agent，执行 Check Policy。", "Agent Online，Host Data/Check Result 可见。", "ACC页面截图", ""],
    ["V-10", "ACC / Agent Based", "Linux", "ACC Agent安装/注册", "安装 ACC Agent，执行 Check Policy。", "Agent Online，Host Data/Check Result 可见。", "ACC页面截图", ""],
    ["V-11", "Network / Peering", "Ubuntu on vnet02 -> vnet03 target", "vnet02↔vnet03 peering Connected；vnet03 NSG 放通测试端口", "从 Ubuntu 运行 ssh/curl/Test-NetConnection 到 vnet03 私有 IP。", "Ubuntu 仅单 NIC，但能通过私网到达 vnet03 目标。", "Connectivity test log / screenshot", ""],
]


RISKS = [
    ["主题", "风险/限制", "处理建议"],
    ["Multi-NIC", "Azure VM 的 NIC 不能跨 VNet；如果 Ubuntu 要同时访问 VNet02 和 VNet03，不能用跨 VNet Multi-NIC。", "Ubuntu 保持单 NIC 放在 vnet02；通过 vnet02↔vnet03 直接 peering 访问 vnet03。"],
    ["Peering", "VNet peering 不是传递式。vnet02 不能靠 vnet01 间接到 vnet03。", "必须直接建立 vnet02↔vnet03 peering，并在 vnet03 NSG 放通 Ubuntu 的测试端口。"],
    ["Azure 私有 IP", "Azure 每个 subnet 保留前几个和最后一个 IP，图中 10.0.0.1/10.0.0.2/10.0.0.3 不建议/不可作为 VM IP。", "实际固定 IP 从 .10/.20 等开始分配。"],
    ["Bastion", "Azure Bastion 需要名为 AzureBastionSubnet 的专用 subnet，且大小通常要求 /26 或更大。", "vnet01 预留 10.0.0.128/26。"],
    ["Windows 11", "Azure Windows client VM 可能受 license / image 可用性约束。", "优先确认订阅/租户资格；若不可用，记录偏差并用 Windows Server 先验证 ACC 流程。"],
    ["Marketplace 设备", "Palo Alto / FortiGate PAYG 会产生软件费用，且部署模板可能要求多个 NIC/subnet。", "先用最低规格、target-only；若模板强制多 subnet，在 vnet01 增加 NVA subnets 并更新资源清单。"],
    ["Service Mapping", "裸 OS 不足以验证 top-down mapping；必须有可访问 application entry point。", "在 Windows 安装 IIS 或在 Linux 安装 Nginx，开放 80/443。"],
    ["Cloud Discovery", "Azure Cloud Discovery 是 API 级资源发现，不等同于 OS 级 credentialed Discovery。", "Cloud Discovery 和 IP Range Discovery 分开验证、分开截图。"],
    ["NSG", "过严 NSG 会导致 Discovery 找不到端口或凭据失败。", "PoC 初期只限制来源为 MID 私有 IP，端口可适度放宽；验收后再收紧。"],
]


TROUBLESHOOTING = [
    ["现象", "优先检查", "处理"],
    ["MID Server Down/无法 Validate", "VM 出站 443、DNS、代理、MID 用户密码、时间同步", "从 VM curl/浏览器访问 instance；确认 ServiceNow 中 MID 用户角色与密码。"],
    ["Windows Credential failed", "Windows 防火墙、WMI/RPC、账号权限、UAC remote restrictions", "先从 MID 测试 135/445/5985；临时开放 Windows Discovery 相关规则。"],
    ["Linux SSH failed", "SSH 22、用户名/密钥、sudo、NSG", "从 MID ssh 测试；确认账号可执行 uname/lsblk/ip 等基础命令。"],
    ["SNMP failed", "设备 SNMP service、community/user、ACL、UDP 161", "从 MID 使用 snmpwalk 或 ServiceNow credential test；确认来源 IP 是 MID 私有 IP。"],
    ["Cloud Discovery 无结果", "Tenant ID、Subscription ID、Client ID/Secret、Reader role、secret 过期", "Azure Portal 重新确认 role assignment；ServiceNow 重新测试 credential。"],
    ["Service Map 为空", "Entry point 可达性、应用服务是否真实监听、目标主机是否已被 Discovery", "从 MID curl 入口；先成功发现 OS CI，再跑 top-down。"],
    ["Bastion 无法创建", "AzureBastionSubnet 名称/大小、Public IP SKU、区域可用性", "确认 subnet 名称完全一致且 /26 以上。"],
    ["VNet 间不通", "Peering 双向状态、NSG、路由表、地址空间重叠、是否直接 peering", "检查 vnet02↔vnet03 都是 Connected；不要依赖 transitive peering；使用 Network Watcher 测试。"],
]


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Aptos"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        st = doc.styles[style_name]
        st.font.name = "Aptos Display"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def set_run_east_asia(run, font_name: str = "Microsoft YaHei") -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Aptos"
            set_run_east_asia(run)
            run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        shade_cell(hdr[i], "1F4E79")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths_cm:
        for row in table.rows:
            for i, width in enumerate(widths_cm):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def create_word() -> None:
    doc = Document()
    set_doc_defaults(doc)
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Azure ServiceNow Discovery / Service Mapping PoC\n环境构筑手册")
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x0B, 0x1F, 0x33)
    set_run_east_asia(r)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"基于用户提供的三张构成/范围图整理｜生成日：{TODAY}").italic = True

    doc.add_paragraph()
    doc.add_paragraph("本手册面向“几乎从 0 开始使用 Azure”的实施者，主线使用 Azure Portal 操作；ServiceNow 部分按 Discovery、Cloud Discovery、Service Mapping、ACC 四条验证线拆分。")
    doc.add_paragraph("重要：本文档不会保存任何真实密码、token、client secret 或 SNMP community。请把凭据存放在 ServiceNow Credential、Azure Key Vault 或组织认可的密码库中。")

    doc.add_heading("1. 结论与构筑方针", level=1)
    add_bullets(doc, [
        "PoC 的核心不是单纯创建 VM，而是构造一组可被 ServiceNow 从不同机制发现的 Azure 目标：IP Range Discovery、Credentialed Discovery、Azure Cloud Discovery、Service Mapping、ACC。",
        "Azure 网络采用 3 个 VNet：vnet01=Server & Network Device Zone，vnet02=MID & Linux Discovery Zone，vnet03=Client & ACC Zone；三者 Full Mesh Peering。",
        "Ubuntu 不再做跨 VNet Multi-NIC，而是单 NIC 放在 vnet02；通过 vnet02↔vnet03 直接 peering 验证它能同时访问两个 VNet。",
        "两个 MID Server：Windows MID 放在 vnet01，Linux MID 放在 vnet02。这样能分别验证 Windows/Linux MID 在不同网络区的可达性和选择策略。",
        "Palo Alto / FortiGate 作为 Discovery 対象のみ（只作为被发现对象），不承担真实流量转发；重点启用管理接口与 SNMP。",
        "按图中范围，AWS / Others Cloud / NW(Others) / traffic-based Mapping 不做。",
    ])

    doc.add_heading("2. 从三张图抽取出的验证范围", level=1)
    add_table(
        doc,
        ["大分类", "中分类", "小分类", "验证对象", "范围", "实施解释"],
        VALIDATION_SCOPE,
        [2.6, 3.2, 3.2, 4.8, 1.8, 6.5],
    )

    doc.add_heading("3. Azure 实作时必须修正/注意的点", level=1)
    add_table(doc, ["主题", "风险/限制", "处理建议"], RISKS, [3.2, 7.2, 7.2])

    doc.add_heading("4. 目标架构", level=1)
    doc.add_paragraph("逻辑架构如下。真实 IP 请以第 5 章的 IP Plan 为准。")
    diagram = (
        "ServiceNow Instance / ITOM Visibility\n"
        "        │ HTTPS 443\n"
        "        ├── Windows MID: vm-win-mid-01 / 10.0.0.10 / vnet01\n"
        "        └── Linux MID:   vm-linux-mid-01 / 172.16.0.10 / vnet02\n\n"
        "vnet01 10.0.0.0/24  Server & Network Device Zone\n"
        "  ├─ Windows Server 2019/2022\n"
        "  ├─ RHEL 9\n"
        "  ├─ Palo Alto VM-Series（Discovery対象のみ）\n"
        "  └─ FortiGate VM（Discovery対象のみ）\n\n"
        "vnet02 172.16.0.0/24  MID & Linux Discovery Zone\n"
        "  ├─ Linux MID\n"
        "  ├─ Ubuntu 22\n"
        "  └─ Ubuntu 22（单 NIC；通过 vnet02↔vnet03 peering 访问 vnet03）\n\n"
        "vnet03 192.168.0.0/24  Client & ACC Zone\n"
        "  └─ Windows 11 / ACC target\n\n"
        "Peering: vnet01 ↔ vnet02, vnet01 ↔ vnet03, vnet02 ↔ vnet03"
    )
    p = doc.add_paragraph()
    run = p.add_run(diagram)
    run.font.name = "Consolas"
    set_run_east_asia(run)
    run.font.size = Pt(9)

    doc.add_heading("5. Azure 资源清单", level=1)
    add_table(doc, ["分类", "资源类型", "资源名", "数量", "配置/范围", "说明"], RESOURCE_PLAN, [2, 3.3, 4.2, 1.4, 4.2, 6.3])

    doc.add_heading("6. IP / Subnet 计划", level=1)
    doc.add_paragraph("注意：Azure 会保留每个 subnet 的前几个和最后一个地址，因此不要按概念图把 .1/.2/.3 分配给 VM。建议从 .10/.20 等地址开始固定。")
    add_table(doc, IP_PLAN[0], IP_PLAN[1:], [3, 2.1, 3.4, 3.1, 4.3, 3, 5.2])

    doc.add_heading("7. NSG 初期规则", level=1)
    doc.add_paragraph("PoC 初期建议以“来源严格、端口适度放宽”为原则：来源只允许 Windows/Linux MID 的私有 IP；待 Discovery 跑通后再按实际端口收紧。")
    add_table(doc, NSG_RULES[0], NSG_RULES[1:], [3, 1.8, 1.6, 4, 3.6, 3.6, 3.6, 1.6, 5.2])

    doc.add_heading("8. 构筑手顺（Azure Portal 主线）", level=1)
    add_table(doc, ["步骤ID", "步骤", "操作内容", "完成条件"], BUILD_STEPS, [1.6, 4.2, 8.5, 5.8])

    doc.add_heading("9. ServiceNow 设置手顺", level=1)
    add_table(doc, SERVICENOW_SETUP[0], SERVICENOW_SETUP[1:], [3, 5.2, 5.2, 4.6, 4])

    doc.add_heading("10. 推荐的详细跟做顺序", level=1)
    add_numbered(doc, [
        "先只创建网络：RG、3个VNet、subnet、NSG、NAT、Bastion、Peering。不要急着创建所有 VM。",
        "创建 Windows MID 和 Linux MID 两台 VM，确认能出站访问 ServiceNow instance。",
        "在 ServiceNow 安装并 Validate 两个 MID。MID 不稳定时不要继续往下做。",
        "先创建 Windows Server 2019 和 Ubuntu 22 两个最小目标，验证 Windows/Linux credentialed Discovery。",
        "再创建 RHEL、Windows Server 2022、Windows 11，逐步扩大范围。",
        "最后用 Ubuntu 22 验证 vnet02↔vnet03 连通性：确认它只有单 NIC，但能访问 vnet03 的测试端口。",
        "部署 Palo Alto/FortiGate，先确认 SNMP 从 MID 可达，再运行网络设备 Discovery。",
        "创建 Azure Service Principal 和 Reader role，运行 Cloud Discovery。Cloud Discovery 与 IP Discovery 分开截图。",
        "安装 IIS/Nginx，创建 Service Mapping Entry Point，运行 top-down mapping。",
        "最后安装 ACC Agent 到 Windows 11 和 Linux，验证 Agent Online、Host Data 和 Check Result。",
        "收集所有证据后停止 VM；如 PoC 完成，删除 Resource Group 以避免持续计费。",
    ])

    doc.add_heading("11. 验收矩阵", level=1)
    add_table(doc, VERIFICATION_MATRIX[0], VERIFICATION_MATRIX[1:], [1.3, 4, 3.2, 4.1, 5.2, 5.2, 3.5, 1.5])

    doc.add_heading("12. 常见故障排查", level=1)
    add_table(doc, ["现象", "优先检查", "处理"], TROUBLESHOOTING, [4.8, 6.5, 6.5])

    doc.add_heading("13. 成本控制与清理", level=1)
    add_bullets(doc, [
        "Palo Alto VM-Series、FortiGate VM、Bastion、NAT Gateway、Public IP、Managed Disk 都可能在 VM 停止后继续产生费用；PoC 不使用时请停止/删除。",
        "每天结束前至少执行：停止 VM、确认 Marketplace 设备不再运行、确认没有多余 Public IP 和磁盘。",
        "PoC 结束后最安全的清理方式：确认没有需要保留的证据文件后，删除 RG-SNOW-Discovery-POC 整个 Resource Group。",
    ])

    doc.add_heading("14. 证据包目录建议", level=1)
    add_table(
        doc,
        ["目录", "内容"],
        [
            ["01_Azure_Build", "RG/VNet/Subnet/NSG/Peering/NAT/Bastion/VM 截图"],
            ["02_MID", "MID Server Up/Validated 截图、MID 版本、运行日志摘要"],
            ["03_Discovery", "Discovery Schedule、Discovery Status、Windows/Linux/Network CI 截图"],
            ["04_Cloud_Discovery", "Azure Credential/Cloud Account/Schedule、Azure CI 截图"],
            ["05_Service_Mapping", "Entry Point、Application Service Map、Pattern/Log 截图"],
            ["06_ACC", "Agent Online、Host Data、Check Result 截图"],
            ["07_Issues", "失败日志、修复记录、范围外/限制说明"],
        ],
        [5.2, 11.5],
    )

    doc.add_heading("15. 官方资料与依据", level=1)
    source_rows = [[s["id"], s["title"], s["url"], s["use"]] for s in SOURCES]
    add_table(doc, ["ID", "官方资料", "URL", "本文用途"], source_rows, [1.3, 5, 8, 5.5])

    doc.add_heading("16. 附录：用户提供的原始参考图", level=1)
    for name, path in SOURCE_IMAGES:
        if path.exists():
            doc.add_heading(name, level=2)
            doc.add_picture(str(path), width=Cm(17))
            doc.add_paragraph(str(path))

    doc.save(DOCX_PATH)


def style_sheet(ws, freeze: str = "A2") -> None:
    ws.freeze_panes = freeze
    ws.sheet_view.showGridLines = False
    thin = Side(style="thin", color="D9E2F3")
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = Border(left=thin, right=thin, top=thin, bottom=thin)
            cell.font = Font(name="Yu Gothic", size=10)
    for cell in ws[1]:
        cell.fill = PatternFill("solid", fgColor="1F4E79")
        cell.font = Font(name="Yu Gothic", size=10, bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def autosize(ws, max_width: int = 60) -> None:
    for col in ws.columns:
        width = 10
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            if cell.value is None:
                continue
            text = str(cell.value)
            if not text:
                continue
            width = max(width, min(max_width, max(len(line) for line in text.splitlines()) + 2))
        ws.column_dimensions[col_letter].width = width


def add_rows(ws, headers: list[str], rows: list[list[str]]) -> None:
    ws.append(headers)
    for row in rows:
        ws.append(row)
    style_sheet(ws)
    autosize(ws)


def create_excel() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Overview"
    overview_rows = [
        ["项目", "内容"],
        ["文档生成日", TODAY],
        ["目的", "Azure 上构筑 ServiceNow Discovery / Service Mapping / ACC PoC 环境"],
        ["默认区域", "Japan East（如客户要求可替换）"],
        ["Resource Group", "RG-SNOW-Discovery-POC"],
        ["关键限制", "Azure VM 多 NIC 不能跨 VNet；如果 Ubuntu 要同时访问 VNet02/VNet03，必须用直接 peering，而不是跨 VNet Multi-NIC；Bastion 需要 AzureBastionSubnet /26；不要使用 Azure subnet 保留 IP。"],
        ["输出文件", str(DOCX_PATH)],
    ]
    for row in overview_rows:
        ws.append(row)
    style_sheet(ws)
    autosize(ws)

    sheets = [
        ("Validation_Scope", ["大分类", "中分类", "小分类", "验证对象", "范围", "实施解释"], VALIDATION_SCOPE),
        ("Azure_Resources", ["分类", "资源类型", "资源名", "数量", "配置/范围", "说明"], RESOURCE_PLAN),
        ("IP_Plan", IP_PLAN[0], IP_PLAN[1:]),
        ("NSG_Rules", NSG_RULES[0], NSG_RULES[1:]),
        ("Build_Checklist", ["Done", "步骤ID", "步骤", "操作内容", "完成条件"], [["", *row] for row in BUILD_STEPS]),
        ("ServiceNow_Setup", SERVICENOW_SETUP[0], SERVICENOW_SETUP[1:]),
        ("Verification_Matrix", VERIFICATION_MATRIX[0], VERIFICATION_MATRIX[1:]),
        ("Risks_Assumptions", ["主题", "风险/限制", "处理建议"], RISKS),
        ("Troubleshooting", ["现象", "优先检查", "处理"], TROUBLESHOOTING),
        ("Sources", ["ID", "官方资料", "URL", "本文用途"], [[s["id"], s["title"], s["url"], s["use"]] for s in SOURCES]),
    ]

    for title, headers, rows in sheets:
        ws2 = wb.create_sheet(title)
        add_rows(ws2, headers, rows)

    # Add source image references to a lightweight sheet. Do not embed all full-size photos to keep the workbook practical.
    ws_img = wb.create_sheet("Input_Images")
    ws_img.append(["名称", "路径", "存在"])
    for name, path in SOURCE_IMAGES:
        ws_img.append([name, str(path), "Yes" if path.exists() else "No"])
    style_sheet(ws_img)
    autosize(ws_img, max_width=90)

    wb.save(XLSX_PATH)


if __name__ == "__main__":
    create_word()
    create_excel()
    print(str(DOCX_PATH).encode("unicode_escape").decode("ascii"))
    print(str(XLSX_PATH).encode("unicode_escape").decode("ascii"))
